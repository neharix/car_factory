import os
import random
import shutil

import docx
import docx2pdf
import pythoncom
import win32com.client
import win32com.client.makepy
import winerror
from django.conf import settings
from django.contrib import messages
from django.contrib.auth import authenticate, get_user_model, login, logout
from django.contrib.auth.decorators import login_required
from django.db.models import Q
from django.shortcuts import redirect, render
from django.views.generic import ListView
from docx.shared import Mm
from win32com.client.dynamic import ERRORS_BAD_CONTEXT

from .forms import CarForm, SampleForm, UserForm
from .models import Car, Letter, Sample, User


def pdf_to_word(input_file, output_file):
    ERRORS_BAD_CONTEXT.append(winerror.E_NOTIMPL)
    src = os.path.abspath(input_file)

    pythoncom.CoInitializeEx(0)
    win32com.client.makepy.GenerateFromTypeLibSpec("Acrobat")
    adobe = win32com.client.DispatchEx("AcroExch.App")
    avDoc = win32com.client.DispatchEx("AcroExch.AVDoc")

    avDoc.Open(src, src)
    pdDoc = avDoc.GetPDDoc()
    jObject = pdDoc.GetJSObject()

    jObject.SaveAs(output_file, "com.adobe.acrobat.docx")
    avDoc.Close(-1)


class CarObject:
    def __init__(self, pk, name, users, vehicle_type, color, year, number, documents, pdf_documents, desc=None,
                 image=None):
        self.pk = pk
        self.desc = desc
        self.name = name
        self.users = users
        self.type = vehicle_type
        self.color = color
        self.year = year
        self.number = number
        self.documents = documents
        self.pdf_documents = pdf_documents
        self.image = image


class UserObject:
    def __init__(self, pk, first_name, last_name, father_name, username, email, phone_number, passport_serie, vehicles,
                 document, pdf_document):
        self.pk = pk
        self.first_name = first_name
        self.last_name = last_name
        self.father_name = father_name
        self.username = username
        self.email = email
        self.phone_number = phone_number
        self.passport_serie = passport_serie
        self.vehicles = vehicles
        self.documents = document
        self.pdf_document = pdf_document


def index(request):
    return render(request, "index.html")

def home(request):
    return render(request, "home.html")

def home_redirecter(request):
    return redirect("home")

def signin(request):
    if request.method == "POST":
        loginusername = request.POST["username"]
        loginpassword = request.POST["password"]

        user = authenticate(username = loginusername,password = loginpassword)
        if user is not None:
            login(request, user)
            # messages.success(request,"Successfully logged in!")
            return redirect("home")
        else:
            messages.error(request,"Ýalňyş maglumat girizildi!")
            return redirect("signin")

    else:
        print("error")
        return render(request,"login.html")

@login_required(login_url="signin")
def signout(request):
    logout(request)
    # messages.success(request,"Successfully logged out!")
    return redirect("signin")


# return HttpResponse("signout")

@login_required(login_url="signin")
def panel(request):
    if request.user.is_staff or request.user.is_superuser:
        cars = Car.objects.all()
        params = {"cars": cars}
        return render(request, "for_staff.html ", params)
    else:
        cars = Car.objects.all()
        params = {"cars": cars}
        return render(request, "for_default.html", params)

class SearchCarResultsListView(ListView):
	model = Car
	template_name = 'car_search_results.html'

	def get_queryset(self): # new
		query = self.request.GET.get('q')
		return Car.objects.filter(
		Q(car_name__icontains=query) | Q(car_desc__icontains=query) | Q(color__color__icontains=query) | Q(vehicle_type__vehicle_type__icontains=query) | Q(vehicle_number__icontains=query) | Q(car_year__year__icontains=query)
		)

class SearchUserResultsListView(ListView):
	model = User
	template_name = 'user_search_results.html'

	def get_queryset(self): # new
		query = self.request.GET.get('q')
		return User.objects.filter(
		Q(first_name__icontains=query) | Q(last_name__icontains=query) | Q(username__icontains=query) | Q(email__icontains=query) | Q(father_name__icontains=query) | Q(passport_serie__icontains=query) | Q(phone_number__icontains=query)
		)

class SearchSampleResultsListView(ListView):
	model = Sample
	template_name = 'sample_search_results.html'

	def get_queryset(self): # new
		query = self.request.GET.get('q')
		return Sample.objects.filter(
		Q(name__icontains=query)
        )


@login_required(login_url="signin")
def add_sample(request):
    if request.user.is_staff or request.user.is_superuser:
        if request.method == "POST":
            form = SampleForm(request.POST, request.FILES)
            if form.is_valid():
                form.save()
                messages.success(request, "Nusga goşuldy!")
            return redirect("panel")
        else:
            context = {"form": SampleForm()}
            return render(request, "add_sample.html", context)
    else:
        return redirect("home")

@login_required(login_url="signin")
def delete_user(request, pk):
    if request.user.is_superuser or request.user.is_staff:
        user = User.objects.get(pk=pk)
        user.delete()
        return redirect("about_users")
    else:
        return redirect("about_users")

@login_required(login_url="signin")
def delete_vehicle(request, pk):
    if request.user.is_superuser or request.user.is_staff:
        vehicle = Car.objects.get(pk=pk)
        vehicle.delete()
        return redirect("about_vehicles")
    else:
        return redirect("about_vehicles")

@login_required(login_url="signin")
def delete_sample(request, pk):
    if request.user.is_superuser or request.user.is_staff:
        sample = Sample.objects.get(pk=pk)
        sample.delete()
        return redirect("samples")
    else:
        return redirect("samples")


@login_required(login_url="signin")
def add_vehicle(request):
    if request.user.is_staff or request.user.is_superuser:
        if request.method == "POST":
            form = CarForm(request.POST, request.FILES)
            if form.is_valid():
                form.save()
                messages.success(request, "Ulag hasaba alyndy!")
                user_list = []
                car = Car.objects.get(car_name=form.cleaned_data["car_name"])

                base_dir = str(settings.BASE_DIR).replace('\\', '/') + '/media/'
                users = Car.objects.get(pk=car.pk).users.all()
                for user in users:
                    user_list.append(f"{user.first_name} {user.last_name}")
                users = ", ".join(user_list)
                doc = docx.Document()
                doc.add_heading("Ulag barada maglumat")
                doc.add_paragraph(f"Ulagyň ady: {car.car_name}")
                doc.add_paragraph(f"Ulagyň barada: {car.car_desc}")
                doc.add_paragraph(f"Ulanyjylar: {users}")
                doc.add_paragraph(f"Ulagyň reňki: {car.color.color}")
                doc.add_paragraph(f"Ulagyň ýyly: {car.car_year.year}")
                doc.add_paragraph(f"Ulagyň nomeri: {car.vehicle_number}")
                doc.add_paragraph(f"Ulagyň görnüşi: {car.vehicle_type.vehicle_type}")
                doc.add_paragraph().add_run().add_picture(f"{base_dir}{car.image}", width=Mm(150), height=None)
                file_name = ''.join(car.car_name.split()) + f"{random.randint(1, 1000000)}.docx"
                doc.save(f"{base_dir}car/documents/{file_name}")
                pythoncom.CoInitializeEx(0)
                pdf = docx2pdf.convert(base_dir + f"car/documents/{file_name}",
                                       base_dir + f"car/pdf_documents/{file_name.split(".")[0]}.pdf")
                car.characteristics_pdf = f"car/pdf_documents/{file_name.split('.')[0]}.pdf"
                car.characteristics_docx = f"car/documents/{file_name}"
                car.save()
            return redirect("panel")
        else:
            context = {"form": CarForm()}
            return render(request, "add_vehicle.html", context)
    else:
        return redirect("home")


@login_required(login_url="signin")
def add_user(request):
    if request.user.is_staff or request.user.is_superuser:
        if request.method == "POST":
            form = UserForm(request.POST, request.FILES)
            print(request.POST)
            print(request.FILES)
            if form.is_valid():
                form.save()
                messages.success(request, "Ulanyjy hasaba alyndy!")
                user = get_user_model().objects.get(username=form.cleaned_data["username"])
                user.set_password(form.cleaned_data['password'])
                file_name = str(user.documents).split("/")[len(str(user.documents).split("/")) - 1]
                file_format = str(user.documents).split(".")[len(str(user.documents).split(".")) - 1]
                base_dir = str(settings.BASE_DIR).replace("\\", "/")
                if file_format == "pdf":
                    pdf_to_word(base_dir + f"/media/user/documents/{file_name}",
                                base_dir + f"/media/user/documents/{file_name.split(".")[0]}.docx")
                    shutil.copy2(base_dir + f"/media/user/documents/{file_name}",
                                 base_dir + f"/media/user/pdf_documents/{file_name}")
                    os.remove(base_dir + f"/media/user/documents/{file_name}")
                    user.documents = f"user/documents/{file_name.split(".")[0]}.docx"
                    user.pdf_documents = f"user/pdf_documents/{file_name}"
                    user.save()
                elif file_format == "docx":
                    pythoncom.CoInitializeEx(0)
                    pdf = docx2pdf.convert(base_dir + f"/media/user/documents/{file_name}",
                                           base_dir + f"/media/user/pdf_documents/{file_name.split(".")[0]}.pdf")
                    user.pdf_documents = f"user/pdf_documents/{file_name.split(".")[0]}.pdf"
                    user.save()
            return redirect("panel")
        else:
            context = {"form": UserForm()}
            return render(request, "add_user.html", context)
    else:
        return redirect("home")


@login_required(login_url="signin")
def about_vehicles(request):
    if request.user.is_staff or request.user.is_superuser:
        cars = Car.objects.all().order_by("pk")
        car_obj = []
        for car in cars:
            user_list = []

            users = Car.objects.get(pk=car.pk).users.all()
            for user in users:
                user_list.append(f"{user.first_name} {user.last_name}")
            car_obj.append(
                CarObject(car.pk, car.car_name, ", ".join(user_list), car.vehicle_type.vehicle_type, car.color.color,
                          car.car_year.year, car.vehicle_number, car.characteristics_docx, car.characteristics_pdf))
        context = {"cars": car_obj}
        return render(request, "vehicle_table.html", context)
    else:
        return redirect("home")


@login_required(login_url="signin")
def about_users(request):
    if request.user.is_staff or request.user.is_superuser:
        users = get_user_model().objects.all().order_by("pk")
        user_obj = []
        for user in users:
            car = "Adyna ulag bellenilmedik"
            for car in Car.objects.all():
                for car_user in car.users.all():
                    if car_user.username == user.username:
                        car = car.car_name
            if user.is_superuser:
                continue
            else:
                user_obj.append(
                    UserObject(user.pk, user.first_name, user.last_name, user.father_name, user.username, user.email,
                               user.phone_number, user.passport_serie, car, user.documents, user.pdf_documents))
        context = {"users": user_obj}
        return render(request, "user_table.html", context)
    else:
        return redirect("home")


@login_required(login_url='signin')
def view_user_pdf(request, user_id):
    user = get_user_model().objects.get(pk=user_id)
    context = {"user": user}
    return render(request, "view_user_pdf.html", context)


@login_required(login_url='signin')
def view_car_pdf(request, car_id):
    car = Car.objects.get(pk=car_id)
    context = {"car": car}
    return render(request, "view_car_pdf.html", context)


@login_required(login_url="signin")
def about_vehicle(request, car_id):
    car = Car.objects.get(pk=car_id)
    user_list = []
    users = Car.objects.get(pk=car.pk).users.all()
    for user in users:
        user_list.append(f"{user.first_name} {user.last_name}")
    car_obj = CarObject(car.pk, car.car_name, ", ".join(user_list), car.vehicle_type.vehicle_type, car.color.color,
                        car.car_year.year, car.vehicle_number, car.characteristics_docx, car.characteristics_pdf,
                        desc=car.car_desc, image=car.image)
    context = {"car": car_obj}
    return render(request, "about_vehicle.html", context)

@login_required(login_url="signin")
def about_user(request, user_id):
    if request.user.is_superuser or request.user.is_staff:
        user = User.objects.get(pk=user_id)
        car = "Adyna ulag bellenilmedik"
        for car in Car.objects.all():
            for car_user in car.users.all():
                if car_user.username == user.username:
                    car = car.car_name
        user_obj = UserObject(user.pk, user.first_name, user.last_name, user.father_name, user.username, user.email, user.phone_number, user.passport_serie, car, user.documents, user.pdf_documents)
        context = {"user": user_obj}
        return render(request, "about_user.html", context)
    else:
        return redirect("home")

@login_required(login_url="signin")
def mailbox(request):
    letters = Letter.objects.all().order_by("-sent")
    return render(request, "mailbox.html", {"letters": letters})

@login_required(login_url="signin")
def read_letter(request, letter_id):
    letter = Letter.objects.get(pk=letter_id)
    letter.is_checked = True
    letter.save()
    user = User.objects.get(username=letter.user)
    return render(request, "letter.html", {"letter": letter, "user": user})

@login_required(login_url="signin")
def send_letter(request):
    if request.method == "POST":
        Letter.objects.create(user=request.user.username, text=request.POST["text"])
        messages.success(request, "Hat ýolbaşçylara ugradyldy!")
        return redirect("panel")
    else:
        return render(request, "send_letter.html")

@login_required(login_url="signin")
def sample_table(request):
    samples = Sample.objects.all()
    context = {"samples": samples}
    return render(request, 'samples.html', context)


@login_required(login_url="signin")
def view_sample_pdf(request, sample_id):
    sample = Sample.objects.get(pk=sample_id)
    context = {"sample": sample}
    return render(request, "view_sample_pdf.html", context)
